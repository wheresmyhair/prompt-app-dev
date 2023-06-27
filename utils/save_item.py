def save_splited_text(splited_text, save_path):
    with open(save_path, 'w', encoding='utf-8') as f:
        for text in splited_text:
            f.write(text + '\n\n\n\n\n')
            
            
def save_docx_from_str(text, save_path):
    import docx
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = docx.Document()
    p = doc.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.bold = False
    run.font.size = Pt(14)
    run.font.name = "仿宋"  # set custom font name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")  # specify the font for East Asian characters

    doc.save(save_path)