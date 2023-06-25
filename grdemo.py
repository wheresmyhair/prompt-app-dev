import gradio as gr
import docx
import base64
from utils.sysfunc import get_app_version, get_model_version, record_init
from os.path import join, basename


def extract_contents(doc_file):
    record_init(cookies)
    doc = docx.Document(doc_file.name)
    doc.save(join(cookies.value['dir'], basename(doc_file.name)))
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text


def clear_content():
    return gr.update(value=None)


def generate(text):
    return "analyzed: " + text


def extract_and_generate(doc_file):
    text = extract_contents(doc_file)
    return {
        output_doc_content: text, 
        output_doc_generate: generate(text), 
        button_retry: gr.update(visible=True), 
        button_flag: gr.update(visible=True),
    }


def save_as_doc(text):
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(join(cookies.value['dir'], f"考察报告_{cookies.value['access_time']}.docx"))
    return {
        area_download: gr.update(visible=True),
        output_file: gr.update(value=doc), 
        button_flag: gr.update(visible=False), 
        button_retry: gr.update(visible=False),
    }

gr_row = lambda: gr.Row()
gr_col = lambda scale: gr.Column(scale=scale)

with gr.Blocks(title="考察报告生成", analytics_enabled=False) as demo:
    gr.HTML(f"<h1 align=\"center\">考察报告生成 {get_app_version()}</h1>")
    cookies = gr.State({})
    with gr_row():
        with gr_col(scale=3):
            with gr.Tab("考察报告"):
                with gr_row():
                    output_doc_generate = gr.components.Textbox(label="生成结果", placeholder='请先上传文档', show_copy_button=True)
                with gr_row():
                    button_flag = gr.Button('保存结果', variant='primary', visible=False)
                    button_retry = gr.Button('重新生成', variant='secondary', visible=False)
            with gr.Tab("上传文档内容预览"):
                output_doc_content = gr.components.Textbox(label="文档内容预览", placeholder='请先上传文档', show_copy_button=True)
        with gr_col(scale=2):
            with gr.Accordion("上传区", open=True) as area_input:
                with gr.Row():
                    input_file = gr.components.File(label="述职报告: doc/docx", file_types=['.doc','.docx'])
                with gr.Row():
                    button_submit = gr.Button("提交", variant="primary")
                    button_clear = gr.Button("清除", variant="secondary")
            with gr.Accordion("下载区", open=True, visible=False) as area_download:
                with gr.Row():
                    output_file = gr.components.File(label="考察报告")
            with gr.Accordion("状态栏", open=True) as area_status:
                with gr.Row():
                    info_model_version = gr.Text(label='当前模型', value=f'{get_model_version()}')
                    info_app_version = gr.Text(label='当前应用版本', value=f'{get_app_version()}')

    button_submit.click(
        fn=extract_and_generate, 
        inputs=input_file, 
        outputs=[output_doc_content,output_doc_generate, button_flag, button_retry], 
        api_name='extractdoc', 
        show_progress=True)
    
    button_clear.click(fn=clear_content, inputs=None, outputs=input_file)
    
    button_flag.click(
        fn=save_as_doc, 
        inputs=output_doc_generate,
        outputs=[area_download, output_file, button_flag, button_retry],
        api_name='savedoc', 
        show_progress=True)
    
    
demo.launch(
    server_name='0.0.0.0',
    server_port=7800,
)